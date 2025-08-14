"""
Unified Multi-Document Summarization Interface
Supports: DOCX, TXT, CSV file formats with advanced processing
"""

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from workflow_prototype import MultiDocumentSummarizer
import re
from collections import Counter
import tempfile
import os
from typing import List, Dict, Optional

# Import libraries with error handling
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class UnifiedFileProcessor:
    """Unified processor for multiple file formats"""
    
    def __init__(self):
        self.supported_formats = ['txt', 'csv']
        if DOCX_AVAILABLE:
            self.supported_formats.append('docx')
    
    def process_files(self, uploaded_files) -> List[Dict[str, str]]:
        """Process uploaded files and extract content"""
        documents = []
        
        if not uploaded_files:
            return documents
        
        for uploaded_file in uploaded_files:
            try:
                file_extension = uploaded_file.name.split('.')[-1].lower()
                
                if file_extension not in self.supported_formats:
                    st.warning(f"❌ Unsupported format: {file_extension}. Supported: {', '.join(self.supported_formats)}")
                    continue
                
                content = self._extract_content(uploaded_file, file_extension)
                
                if content and content.strip():
                    documents.append({
                        'filename': uploaded_file.name,
                        'content': content,
                        'format': file_extension.upper(),
                        'size': len(content),
                        'word_count': len(content.split())
                    })
                    st.success(f"✅ Processed: {uploaded_file.name}")
                else:
                    st.warning(f"⚠️ No content found in: {uploaded_file.name}")
                    
            except Exception as e:
                st.error(f"❌ Error processing {uploaded_file.name}: {str(e)}")
                
        return documents
    
    def _extract_content(self, file, file_extension: str) -> str:
        """Extract content based on file type"""
        if file_extension == 'txt':
            return self._process_txt(file)
        elif file_extension == 'docx':
            return self._process_docx(file)
        elif file_extension == 'csv':
            return self._process_csv(file)
        return ""
    
    def _process_txt(self, file) -> str:
        """Process TXT files with encoding detection"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                file.seek(0)
                content = file.read().decode(encoding)
                if content.strip():
                    return content
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        # Fallback with error handling
        file.seek(0)
        return file.read().decode('utf-8', errors='ignore')
    
    def _process_docx(self, file) -> str:
        """Process DOCX files"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not installed. Run: pip install python-docx")
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(file.read())
            tmp_file.flush()
            
            try:
                doc = Document(tmp_file.name)
                content = ""
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content += paragraph.text + "\n"
                
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                content += cell.text + " "
                        content += "\n"
                
                return content.strip()
                
            finally:
                os.unlink(tmp_file.name)
                file.seek(0)
    
    def _process_csv(self, file) -> str:
        """Process CSV files"""
        try:
            file.seek(0)
            df = pd.read_csv(file)
            
            content = ""
            for column in df.columns:
                column_data = df[column].dropna().astype(str)
                if not column_data.empty:
                    content += f"{column}: {' '.join(column_data.values)}\n"
            
            return content.strip()
            
        except Exception:
            # Fallback: treat as text
            file.seek(0)
            return file.read().decode('utf-8', errors='ignore')
        finally:
            file.seek(0)

class UnifiedSummarizer(MultiDocumentSummarizer):
    """Unified summarizer with detailed step-by-step visualization"""
    
    def __init__(self):
        super().__init__()
        self.file_processor = UnifiedFileProcessor()
    
    def process_documents(self, documents: List[Dict[str, str]], reduction_factor: float = 0.1):
        """Complete document processing workflow with detailed visualization"""
        
        if not documents:
            st.error("No documents to process!")
            return None
        
        # Step 1: Detailed Preprocessing Visualization
        all_processed_docs = self.visualize_detailed_preprocessing(documents)
        
        if not any(all_processed_docs):
            st.error("No valid sentences after preprocessing!")
            return None
        
        st.markdown("---")
        
        # Step 2: Word Frequency Analysis with Enhanced Visualization
        word_scores = self.visualize_word_frequency_analysis(all_processed_docs, documents)
        
        if not word_scores:
            st.error("No words found for analysis!")
            return None
        
        st.markdown("---")
        
        # Step 3: Sentence Scoring with Line Graphs
        all_sentence_data = self.visualize_sentence_scoring_detailed(all_processed_docs, documents, word_scores)
        
        st.markdown("---")
        
        # Step 4: Summary Generation with Selection Visualization
        final_summary = self.visualize_summary_generation_detailed(all_sentence_data, reduction_factor)
        
        if final_summary:
            return final_summary
        else:
            st.error("Failed to generate summary!")
            return None
    
    def visualize_detailed_preprocessing(self, documents: List[Dict[str, str]]):
        """Detailed preprocessing visualization showing each step"""
        
        st.subheader("📝 Step 1: Detailed Document Preprocessing")
        
        all_processed_docs = []
        
        # Create tabs for each document if multiple documents
        if len(documents) > 1:
            doc_tabs = st.tabs([f"📄 {doc['filename'][:15]}{'...' if len(doc['filename']) > 15 else ''}" for doc in documents])
        else:
            doc_tabs = [st.container()]
        
        for i, (doc, tab) in enumerate(zip(documents, doc_tabs)):
            with tab:
                st.write(f"**Document: {doc['filename']} ({doc['format']})**")
                
                # Original text preview
                st.write("**Original Text:**")
                preview_text = doc['content'][:400] + "..." if len(doc['content']) > 400 else doc['content']
                st.text_area("", preview_text, height=120, disabled=True, key=f"original_{i}")
                
                # Step-by-step preprocessing
                processed_sentences = self._process_document_with_steps(doc['content'], i)
                all_processed_docs.append(processed_sentences)
        
        return all_processed_docs
    
    def _process_document_with_steps(self, text: str, doc_index: int):
        """Process document with detailed step visualization"""
        
        # Step 1a: Sentence splitting
        sentences = self._split_into_sentences(text.lower())
        st.write(f"**Step 1a: Sentence Splitting ({len(sentences)} sentences)**")
        
        with st.expander(f"View split sentences"):
            for i, sentence in enumerate(sentences[:5], 1):
                st.write(f"S{i}: {sentence}")
            if len(sentences) > 5:
                st.write(f"... and {len(sentences) - 5} more sentences")
        
        # Step 1b: Stop words removal
        st.write("**Step 1b: Stop Words Removal**")
        sentences_after_stopwords = []
        
        for sentence in sentences:
            words = re.findall(r'\b[a-z]+\b', sentence)
            filtered_words = [word for word in words if word not in self.stop_words]
            if filtered_words:
                sentences_after_stopwords.append(' '.join(filtered_words))
        
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Before stop word removal", len(sentences))
        with col2:
            st.metric("After stop word removal", len(sentences_after_stopwords))
        
        with st.expander("View sentences after stop word removal"):
            for i, sentence in enumerate(sentences_after_stopwords[:3], 1):
                st.write(f"S{i}: {sentence}")
            if len(sentences_after_stopwords) > 3:
                st.write(f"... and {len(sentences_after_stopwords) - 3} more")
        
        # Step 1c: Cue words removal
        st.write("**Step 1c: Cue Words Removal**")
        sentences_after_cuewords = []
        
        for sentence in sentences_after_stopwords:
            words = sentence.split()
            filtered_words = [word for word in words if word not in self.cue_words]
            if filtered_words:
                sentences_after_cuewords.append(' '.join(filtered_words))
        
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Before cue word removal", len(sentences_after_stopwords))
        with col2:
            st.metric("After cue word removal", len(sentences_after_cuewords))
        
        with st.expander("View sentences after cue word removal"):
            for i, sentence in enumerate(sentences_after_cuewords[:3], 1):
                st.write(f"S{i}: {sentence}")
            if len(sentences_after_cuewords) > 3:
                st.write(f"... and {len(sentences_after_cuewords) - 3} more")
        
        # Step 1d: Dictionary words removal
        st.write("**Step 1d: Basic Dictionary Words Removal**")
        final_processed_sentences = []
        
        for sentence in sentences_after_cuewords:
            words = sentence.split()
            filtered_words = [
                word for word in words 
                if word not in self.basic_dictionary_words and len(word) > 1
            ]
            if filtered_words:
                final_processed_sentences.append(' '.join(filtered_words))
        
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Before dictionary word removal", len(sentences_after_cuewords))
        with col2:
            st.metric("Final processed sentences", len(final_processed_sentences))
        
        with st.expander("View final processed sentences"):
            for i, sentence in enumerate(final_processed_sentences[:5], 1):
                st.write(f"S{i}: {sentence}")
            if len(final_processed_sentences) > 5:
                st.write(f"... and {len(final_processed_sentences) - 5} more")
        
        return final_processed_sentences
    
    def visualize_word_frequency_analysis(self, all_processed_docs: List[List[str]], documents: List[Dict[str, str]]):
        """Enhanced word frequency analysis with detailed visualization"""
        
        st.subheader("📊 Step 2: Word Frequency Analysis")
        
        # Calculate word frequencies
        all_word_freq = Counter()
        doc_word_freqs = []
        
        for processed_sentences in all_processed_docs:
            doc_freq = Counter()
            for sentence in processed_sentences:
                words = sentence.split()
                doc_freq.update(words)
                all_word_freq.update(words)
            doc_word_freqs.append(doc_freq)
        
        if not all_word_freq:
            return {}
        
        # Overall frequency analysis
        st.write("**Overall Word Frequency Analysis:**")
        top_words = all_word_freq.most_common(20)
        words_df = pd.DataFrame(top_words, columns=['Word', 'Frequency'])
        
        # Interactive bar chart
        fig = px.bar(words_df, x='Word', y='Frequency', 
                    title='Top 20 Most Frequent Words Across All Documents',
                    color='Frequency',
                    color_continuous_scale='viridis')
        fig.update_xaxes(tickangle=45)
        st.plotly_chart(fig, use_container_width=True)
        
        # Per-document analysis if multiple documents
        if len(documents) > 1:
            st.write("**Per-Document Word Distribution:**")
            
            # Create frequency comparison
            freq_comparison = []
            for word, total_freq in top_words[:10]:
                row = {'Word': word, 'Total': total_freq}
                for i, (doc_freq, doc) in enumerate(zip(doc_word_freqs, documents)):
                    row[f"Doc_{i+1}"] = doc_freq.get(word, 0)
                freq_comparison.append(row)
            
            comparison_df = pd.DataFrame(freq_comparison)
            
            # Stacked bar chart for document comparison
            fig_stacked = go.Figure()
            colors = px.colors.qualitative.Set3
            
            for i, doc in enumerate(documents):
                fig_stacked.add_trace(go.Bar(
                    name=f'{doc["filename"][:15]}...' if len(doc["filename"]) > 15 else doc["filename"],
                    x=comparison_df['Word'],
                    y=comparison_df[f'Doc_{i+1}'],
                    marker_color=colors[i % len(colors)]
                ))
            
            fig_stacked.update_layout(
                title='Word Frequency Distribution Across Documents',
                barmode='stack',
                xaxis_title='Words',
                yaxis_title='Frequency',
                xaxis_tickangle=45
            )
            st.plotly_chart(fig_stacked, use_container_width=True)
        
        # Word scores calculation and display
        max_freq = max(all_word_freq.values())
        word_scores = {word: freq / max_freq for word, freq in all_word_freq.items()}
        
        # Display word scores table
        score_df = pd.DataFrame(list(word_scores.items()), columns=['Word', 'Score'])
        score_df = score_df.sort_values('Score', ascending=False).head(15)
        
        st.write("**Word Scores Table (Top 15):**")
        st.dataframe(score_df, use_container_width=True)
        
        return word_scores
    
    def visualize_sentence_scoring_detailed(self, all_processed_docs: List[List[str]], 
                                          documents: List[Dict[str, str]], word_scores: dict):
        """Detailed sentence scoring with line graphs and visualizations"""
        
        st.subheader("🎯 Step 3: Sentence Scoring Analysis")
        
        all_sentence_data = []
        all_lengths = []
        doc_sentence_counts = []
        
        # Process each document
        for doc_idx, (processed_sentences, doc_info) in enumerate(zip(all_processed_docs, documents)):
            for sent_idx, sentence in enumerate(processed_sentences, 1):
                words = sentence.split()
                primary_score = sum(word_scores.get(word, 0) for word in words)
                length = len(words)
                all_lengths.append(length)
                
                all_sentence_data.append({
                    'Document': doc_info['filename'],
                    'Doc_ID': doc_idx + 1,
                    'Sentence_ID': f'D{doc_idx+1}S{sent_idx}',
                    'Sentence': sentence,
                    'Length': length,
                    'Primary_Score': round(primary_score, 2)
                })
            
            doc_sentence_counts.append(len(processed_sentences))
        
        # Calculate average length and final scores
        avg_length = sum(all_lengths) / len(all_lengths) if all_lengths else 1
        
        for data in all_sentence_data:
            final_score = data['Primary_Score'] * (avg_length / data['Length']) if data['Length'] > 0 else 0
            data['Final_Score'] = round(final_score, 2)
        
        # Display scoring formula and statistics
        st.write(f"**Average Sentence Length Across All Documents:** {avg_length:.2f}")
        st.write("**Scoring Formula:** Final Score = Primary Score × (Average Length / Current Length)")
        
        # Document statistics
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total Documents", len(documents))
        with col2:
            st.metric("Total Sentences", len(all_sentence_data))
        with col3:
            st.metric("Avg Sentences/Doc", f"{sum(doc_sentence_counts) / len(doc_sentence_counts):.1f}")
        
        # Line graph visualization for sentence scores
        st.write("**Sentence Scores Visualization:**")
        
        if len(documents) == 1:
            # Single document - simple line chart
            fig = go.Figure()
            
            fig.add_trace(go.Scatter(
                x=[d['Sentence_ID'] for d in all_sentence_data],
                y=[d['Primary_Score'] for d in all_sentence_data],
                mode='lines+markers',
                name='Primary Score',
                line=dict(color='blue', width=2),
                marker=dict(size=6)
            ))
            
            fig.add_trace(go.Scatter(
                x=[d['Sentence_ID'] for d in all_sentence_data],
                y=[d['Final_Score'] for d in all_sentence_data],
                mode='lines+markers',
                name='Final Score',
                line=dict(color='red', width=2),
                marker=dict(size=6)
            ))
            
            fig.update_layout(
                title='Sentence Scores Evolution',
                xaxis_title='Sentences',
                yaxis_title='Score',
                xaxis_tickangle=45,
                hovermode='x unified'
            )
            
        else:
            # Multiple documents - grouped line chart
            fig = go.Figure()
            colors = px.colors.qualitative.Set3
            
            for doc_idx, doc in enumerate(documents):
                doc_data = [d for d in all_sentence_data if d['Doc_ID'] == doc_idx + 1]
                doc_name = doc['filename'][:15] + '...' if len(doc['filename']) > 15 else doc['filename']
                
                # Primary scores
                fig.add_trace(go.Scatter(
                    x=[d['Sentence_ID'] for d in doc_data],
                    y=[d['Primary_Score'] for d in doc_data],
                    mode='lines+markers',
                    name=f'{doc_name} - Primary',
                    line=dict(color=colors[doc_idx % len(colors)], width=2, dash='solid'),
                    marker=dict(size=5, symbol='circle'),
                    legendgroup=f'doc_{doc_idx}'
                ))
                
                # Final scores
                fig.add_trace(go.Scatter(
                    x=[d['Sentence_ID'] for d in doc_data],
                    y=[d['Final_Score'] for d in doc_data],
                    mode='lines+markers',
                    name=f'{doc_name} - Final',
                    line=dict(color=colors[doc_idx % len(colors)], width=2, dash='dash'),
                    marker=dict(size=5, symbol='diamond'),
                    legendgroup=f'doc_{doc_idx}'
                ))
            
            fig.update_layout(
                title='Sentence Scores Comparison Across Documents',
                xaxis_title='Sentences',
                yaxis_title='Score',
                xaxis_tickangle=45,
                hovermode='x unified'
            )
        
        st.plotly_chart(fig, use_container_width=True)
        
        # Score distribution histogram
        st.write("**Score Distribution Analysis:**")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Primary score distribution
            fig_hist1 = px.histogram(
                [d['Primary_Score'] for d in all_sentence_data],
                nbins=20,
                title='Primary Score Distribution',
                labels={'value': 'Primary Score', 'count': 'Number of Sentences'}
            )
            st.plotly_chart(fig_hist1, use_container_width=True)
        
        with col2:
            # Final score distribution
            fig_hist2 = px.histogram(
                [d['Final_Score'] for d in all_sentence_data],
                nbins=20,
                title='Final Score Distribution',
                labels={'value': 'Final Score', 'count': 'Number of Sentences'}
            )
            st.plotly_chart(fig_hist2, use_container_width=True)
        
        # Top scoring sentences table
        st.write("**Top 15 Sentences by Final Score:**")
        scoring_df = pd.DataFrame(all_sentence_data)
        top_sentences = scoring_df.nlargest(15, 'Final_Score')
        
        display_columns = ['Document', 'Sentence_ID', 'Length', 'Primary_Score', 'Final_Score']
        st.dataframe(top_sentences[display_columns], use_container_width=True)
        
        return all_sentence_data
    
    def visualize_summary_generation_detailed(self, all_sentence_data: List[dict], reduction_factor: float):
        """Detailed summary generation with selection visualization"""
        
        st.subheader("📄 Step 4: Summary Generation Process")
        
        # Calculate threshold
        final_scores = [data['Final_Score'] for data in all_sentence_data]
        threshold = sum(final_scores) / len(final_scores)
        
        st.write(f"**Threshold Calculation:**")
        st.write(f"Threshold = Sum of Final Scores / Number of Sentences = {threshold:.3f}")
        
        # Select sentences above threshold
        above_threshold = [d for d in all_sentence_data if d['Final_Score'] >= threshold]
        above_threshold.sort(key=lambda x: x['Final_Score'], reverse=True)
        
        # Apply reduction factor
        max_sentences = max(1, int(len(all_sentence_data) * reduction_factor))
        final_selected = above_threshold[:max_sentences]
        
        # Display selection statistics
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Sentences", len(all_sentence_data))
        with col2:
            st.metric("Above Threshold", len(above_threshold))
        with col3:
            st.metric("Reduction Factor", f"{reduction_factor*100}%")
        with col4:
            st.metric("Final Selection", len(final_selected))
        
        # Threshold visualization
        st.write("**Sentence Selection Visualization:**")
        
        # Create selection status for visualization
        selection_data = []
        for data in all_sentence_data:
            if any(s['Sentence_ID'] == data['Sentence_ID'] for s in final_selected):
                status = 'Selected for Summary'
                color = 'green'
            elif data['Final_Score'] >= threshold:
                status = 'Above Threshold'
                color = 'orange'
            else:
                status = 'Below Threshold'
                color = 'red'
            
            selection_data.append({
                'Sentence_ID': data['Sentence_ID'],
                'Final_Score': data['Final_Score'],
                'Status': status,
                'Color': color
            })
        
        # Selection scatter plot
        fig_selection = go.Figure()
        
        for status in ['Selected for Summary', 'Above Threshold', 'Below Threshold']:
            status_data = [d for d in selection_data if d['Status'] == status]
            if status_data:
                color_map = {'Selected for Summary': 'green', 'Above Threshold': 'orange', 'Below Threshold': 'red'}
                
                fig_selection.add_trace(go.Scatter(
                    x=[d['Sentence_ID'] for d in status_data],
                    y=[d['Final_Score'] for d in status_data],
                    mode='markers',
                    name=status,
                    marker=dict(
                        color=color_map[status],
                        size=8,
                        symbol='circle' if status == 'Selected for Summary' else 'circle-open'
                    )
                ))
        
        # Add threshold line
        fig_selection.add_hline(
            y=threshold, 
            line_dash="dash", 
            line_color="black",
            annotation_text=f"Threshold: {threshold:.3f}"
        )
        
        fig_selection.update_layout(
            title='Sentence Selection Process',
            xaxis_title='Sentences',
            yaxis_title='Final Score',
            xaxis_tickangle=45,
            showlegend=True
        )
        
        st.plotly_chart(fig_selection, use_container_width=True)
        
        # Selected sentences details
        st.write("**Selected Sentences for Summary:**")
        if final_selected:
            for i, sent_data in enumerate(final_selected, 1):
                with st.expander(f"Sentence {i} (Score: {sent_data['Final_Score']:.3f})"):
                    st.write(f"**Document:** {sent_data['Document']}")
                    st.write(f"**ID:** {sent_data['Sentence_ID']}")
                    st.write(f"**Content:** {sent_data['Sentence']}")
        else:
            st.warning("No sentences selected for summary!")
        
        # Generate final summary
        if final_selected:
            summary_sentences = [data['Sentence'] for data in final_selected]
            final_summary = '. '.join(summary_sentences) + '.'
            
            st.write("**Generated Summary:**")
            st.success(final_summary)
            
            return final_summary
        else:
            return None

def main():
    """Main application"""
    
    st.set_page_config(
        page_title="Multi-Document Summarization",
        page_icon="📝",
        layout="wide"
    )
    
    st.title("📝 Multi-Document Summarization Tool")
    st.markdown("### Unified interface for DOCX, TXT, and CSV file processing")
    
    processor = UnifiedFileProcessor()
    
    # Sidebar configuration
    st.sidebar.header("⚙️ Settings")
    reduction_factor = st.sidebar.slider(
        "Summary Reduction Factor", 
        min_value=0.05, 
        max_value=0.5, 
        value=0.1, 
        step=0.05,
        help="Percentage of sentences to include in summary"
    )
    
    # File format status
    st.sidebar.header("📚 Format Support")
    st.sidebar.success("✅ TXT files")
    st.sidebar.success("✅ CSV files")
    
    if DOCX_AVAILABLE:
        st.sidebar.success("✅ DOCX files")
    else:
        st.sidebar.warning("⚠️ DOCX support unavailable")
        st.sidebar.code("pip install python-docx")
    
    # Input method selection
    st.sidebar.header("📁 Input Method")
    input_method = st.sidebar.selectbox(
        "Choose input method:",
        ["Upload Files", "Manual Text Input", "Sample Documents"]
    )
    
    documents = []
    
    if input_method == "Upload Files":
        st.header("📁 File Upload")
        st.info(f"**Supported formats:** {', '.join(processor.supported_formats).upper()}")
        
        uploaded_files = st.file_uploader(
            "Choose files",
            type=processor.supported_formats,
            accept_multiple_files=True,
            help="Upload multiple files for processing"
        )
        
        if uploaded_files:
            with st.spinner("Processing files..."):
                documents = processor.process_files(uploaded_files)
    
    elif input_method == "Manual Text Input":
        st.header("📝 Text Input")
        
        num_docs = st.number_input("Number of documents", min_value=1, max_value=5, value=2)
        
        for i in range(num_docs):
            text_input = st.text_area(
                f"Document {i+1}:",
                height=120,
                placeholder="Enter your text here...",
                key=f"text_{i}"
            )
            
            if text_input.strip():
                documents.append({
                    'filename': f'Document_{i+1}.txt',
                    'content': text_input,
                    'format': 'TXT',
                    'size': len(text_input),
                    'word_count': len(text_input.split())
                })
    
    else:  # Sample Documents
        st.header("📋 Sample Documents")
        
        sample_docs = [
            {
                'filename': 'Legal_Article_35A.txt',
                'content': """Articles 35A lay down that only permanent resident of JK shall own immovable property in the state, or get govt jobs or scholarships. It empowers the state for bestowing special rights and privileges to the people. Article 35A has been protecting culture of indigenous people of JK and Ladakh, and their rights to own the land in the state.""",
                'format': 'TXT'
            },
            {
                'filename': 'Legal_Article_370.txt',
                'content': """Introducing Articles 370. Article 370 accords special rights and privileges to J&K citizens, including an exemption from constitutional provisions governing other States. It also empowers the state legislative assembly to frame any law without attracting a legal challenge. Because of the allowance by the Articles 370, the state of J K has its own Constitution.""",
                'format': 'TXT'
            },
            {
                'filename': 'Technology_AI.txt',
                'content': """Artificial Intelligence is revolutionizing the way we work and live. Machine learning algorithms can process vast amounts of data to identify patterns and make predictions. Deep learning neural networks are particularly effective for image recognition and natural language processing tasks. Companies are investing heavily in AI research and development.""",
                'format': 'TXT'
            }
        ]
        
        # Calculate additional fields
        for doc in sample_docs:
            doc['size'] = len(doc['content'])
            doc['word_count'] = len(doc['content'].split())
        
        documents = sample_docs
        st.success(f"✅ Loaded {len(documents)} sample documents")
        
        # Preview documents
        for doc in documents:
            with st.expander(f"📄 {doc['filename']}"):
                st.text(doc['content'])
    
    # Process documents
    if st.button("🚀 Process Documents", type="primary", disabled=len(documents) == 0):
        if documents:
            with st.spinner("Processing documents through summarization workflow..."):
                summarizer = UnifiedSummarizer()
                
                try:
                    final_summary = summarizer.process_documents(documents, reduction_factor)
                    
                    if final_summary:
                        # Results section
                        st.markdown("---")
                        st.header("🎉 Processing Complete")
                        
                        # Statistics
                        total_words = sum(doc['word_count'] for doc in documents)
                        summary_words = len(final_summary.split())
                        reduction = (1 - summary_words / total_words) * 100 if total_words > 0 else 0
                        
                        col1, col2, col3, col4 = st.columns(4)
                        with col1:
                            st.metric("Documents", len(documents))
                        with col2:
                            st.metric("Original Words", f"{total_words:,}")
                        with col3:
                            st.metric("Summary Words", f"{summary_words:,}")
                        with col4:
                            st.metric("Reduction", f"{reduction:.1f}%")
                        
                        # Download button
                        st.download_button(
                            label="📥 Download Summary",
                            data=final_summary,
                            file_name="document_summary.txt",
                            mime="text/plain"
                        )
                
                except Exception as e:
                    st.error(f"❌ Processing error: {str(e)}")
                    st.info("Please check your input and try again.")
        else:
            st.warning("Please provide documents to process!")
    
    # Help section
    st.markdown("---")
    st.header("ℹ️ About This Tool")
    
    with st.expander("📖 How It Works"):
        st.write("""
        **Multi-Document Summarization Process:**
        1. **Preprocessing**: Clean and prepare text from multiple documents
        2. **Word Analysis**: Calculate word frequencies across all documents
        3. **Sentence Scoring**: Score sentences based on word importance and length
        4. **Summary Generation**: Select top sentences using threshold and reduction factor
        
        **Supported Formats:**
        - 📄 **TXT**: Plain text files with automatic encoding detection
        - 📊 **CSV**: Structured data converted to readable text
        - 📝 **DOCX**: Microsoft Word documents (requires python-docx)
        """)
    
    with st.expander("🔧 Technical Information"):
        st.write("""
        **Scoring Algorithm:**
        - Primary Score = Sum of normalized word frequencies in sentence
        - Final Score = Primary Score × (Average Length / Current Length)
        - Threshold = Average of all final scores
        
        **Installation Requirements:**
        ```bash
        pip install streamlit pandas plotly python-docx
        ```
        """)

if __name__ == "__main__":
    main()